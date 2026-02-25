"""
完整的文件解析工具
支持PDF、DOCX、TXT、HTML等多种格式的解析
"""
import os
import re
import magic
from pathlib import Path
from typing import Dict, List, Optional, Union, Any
from datetime import datetime
import logging

# PDF解析
try:
    import fitz  # PyMuPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# DOCX解析
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# HTML解析
try:
    from bs4 import BeautifulSoup
    HTML_AVAILABLE = True
except ImportError:
    HTML_AVAILABLE = False

# 设置日志
logger = logging.getLogger(__name__)

class FileParser:
    """文件解析器"""
    
    # 支持的文件类型
    SUPPORTED_TYPES = {
        '.pdf': 'pdf',
        '.docx': 'docx', 
        '.doc': 'docx',  # 旧版Word文档，尝试用docx解析
        '.txt': 'txt',
        '.text': 'txt',
        '.html': 'html',
        '.htm': 'html'
    }
    
    # MIME类型映射
    MIME_TYPE_MAP = {
        'application/pdf': 'pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        'application/msword': 'docx',  # 旧版Word
        'text/plain': 'txt',
        'text/html': 'html',
        'application/octet-stream': None  # 需要根据文件内容判断
    }
    
    def __init__(self):
        self.max_file_size = 50 * 1024 * 1024  # 50MB
        self.max_text_length = 100000  # 100K字符
        
    def get_file_type(self, file_path: str) -> Optional[str]:
        """
        获取文件类型
        
        Args:
            file_path: 文件路径
            
        Returns:
            文件类型字符串或None
        """
        try:
            path = Path(file_path)
            if not path.exists():
                return None
                
            # 首先尝试通过扩展名判断
            suffix = path.suffix.lower()
            if suffix in self.SUPPORTED_TYPES:
                return self.SUPPORTED_TYPES[suffix]
            
            # 如果扩展名不支持，尝试通过MIME类型判断
            try:
                mime_type = magic.from_file(str(file_path), mime=True)
                file_type = self.MIME_TYPE_MAP.get(mime_type)
                if file_type:
                    return file_type
            except Exception as e:
                logger.warning(f"MIME类型检测失败: {e}")
            
            # 最后尝试通过文件内容判断
            return self._detect_by_content(file_path)
            
        except Exception as e:
            logger.error(f"文件类型检测失败: {e}")
            return None
    
    def _detect_by_content(self, file_path: str) -> Optional[str]:
        """通过文件内容判断文件类型"""
        try:
            with open(file_path, 'rb') as f:
                header = f.read(1024)
                
            # PDF文件头
            if header.startswith(b'%PDF-'):
                return 'pdf'
            
            # ZIP文件头（DOCX是ZIP格式）
            if header.startswith(b'PK\x03\x04'):
                return 'docx'
            
            # HTML文件头
            if b'<html' in header.lower() or b'<!doctype' in header.lower():
                return 'html'
            
            # 纯文本文件（尝试UTF-8解码）
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    f.read(100)
                return 'txt'
            except UnicodeDecodeError:
                pass
                
        except Exception as e:
            logger.warning(f"内容检测失败: {e}")
        
        return None
    
    def parse_file(self, file_path: str, file_type: Optional[str] = None) -> Dict[str, Any]:
        """
        解析文件内容
        
        Args:
            file_path: 文件路径
            file_type: 文件类型，如果为None则自动检测
            
        Returns:
            解析结果字典
        """
        try:
            # 检查文件是否存在
            if not os.path.exists(file_path):
                return {
                    "success": False,
                    "error": "文件不存在",
                    "file_type": None,
                    "content": None
                }
            
            # 检查文件大小
            file_size = os.path.getsize(file_path)
            if file_size > self.max_file_size:
                return {
                    "success": False,
                    "error": f"文件大小超过限制（最大{self.max_file_size // 1024 // 1024}MB）",
                    "file_type": None,
                    "content": None
                }
            
            # 自动检测文件类型
            if file_type is None:
                file_type = self.get_file_type(file_path)
                if file_type is None:
                    return {
                        "success": False,
                        "error": "无法识别文件类型",
                        "file_type": None,
                        "content": None
                    }
            
            # 根据文件类型选择解析方法
            if file_type == 'pdf':
                return self._parse_pdf(file_path)
            elif file_type == 'docx':
                return self._parse_docx(file_path)
            elif file_type == 'txt':
                return self._parse_txt(file_path)
            elif file_type == 'html':
                return self._parse_html(file_path)
            else:
                return {
                    "success": False,
                    "error": f"不支持的文件类型: {file_type}",
                    "file_type": file_type,
                    "content": None
                }
                
        except Exception as e:
            logger.error(f"文件解析失败: {e}")
            return {
                "success": False,
                "error": f"文件解析失败: {str(e)}",
                "file_type": file_type,
                "content": None
            }
    
    def _parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """解析PDF文件"""
        if not PDF_AVAILABLE:
            return {
                "success": False,
                "error": "PDF解析库未安装，请安装PyMuPDF",
                "file_type": "pdf",
                "content": None
            }
        
        try:
            with fitz.open(file_path) as doc:
                text = ""
                pages = []
                
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    page_text = page.get_text()
                    text += page_text + "\n"
                    pages.append({
                        "page_number": page_num + 1,
                        "text": page_text,
                        "char_count": len(page_text)
                    })
                
                # 清理文本
                text = self._clean_text(text)
                
                return {
                    "success": True,
                    "error": None,
                    "file_type": "pdf",
                    "content": text,
                    "metadata": {
                        "total_pages": len(doc),
                        "pages": pages,
                        "char_count": len(text)
                    }
                }
                
        except Exception as e:
            return {
                "success": False,
                "error": f"PDF解析失败: {str(e)}",
                "file_type": "pdf",
                "content": None
            }
    
    def _parse_docx(self, file_path: str) -> Dict[str, Any]:
        """解析DOCX文件"""
        if not DOCX_AVAILABLE:
            return {
                "success": False,
                "error": "DOCX解析库未安装，请安装python-docx",
                "file_type": "docx",
                "content": None
            }
        
        try:
            doc = Document(file_path)
            text = ""
            paragraphs = []
            
            # 提取段落文本
            for para in doc.paragraphs:
                para_text = para.text.strip()
                if para_text:
                    text += para_text + "\n"
                    paragraphs.append({
                        "text": para_text,
                        "style": para.style.name if para.style else "Normal"
                    })
            
            # 提取表格文本
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        row_data.append(cell_text)
                    table_data.append(row_data)
                tables.append(table_data)
                
                # 将表格转换为文本
                for row in table_data:
                    text += " | ".join(row) + "\n"
            
            # 清理文本
            text = self._clean_text(text)
            
            return {
                "success": True,
                "error": None,
                "file_type": "docx",
                "content": text,
                "metadata": {
                    "paragraphs": paragraphs,
                    "tables": tables,
                    "char_count": len(text)
                }
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"DOCX解析失败: {str(e)}",
                "file_type": "docx",
                "content": None
            }
    
    def _parse_txt(self, file_path: str) -> Dict[str, Any]:
        """解析TXT文件"""
        try:
            # 尝试多种编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'latin-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    
                    # 清理文本
                    text = self._clean_text(text)
                    
                    return {
                        "success": True,
                        "error": None,
                        "file_type": "txt",
                        "content": text,
                        "metadata": {
                            "encoding": encoding,
                            "char_count": len(text)
                        }
                    }
                    
                except UnicodeDecodeError:
                    continue
            
            return {
                "success": False,
                "error": "无法识别文件编码",
                "file_type": "txt",
                "content": None
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"TXT解析失败: {str(e)}",
                "file_type": "txt",
                "content": None
            }
    
    def _parse_html(self, file_path: str) -> Dict[str, Any]:
        """解析HTML文件"""
        if not HTML_AVAILABLE:
            return {
                "success": False,
                "error": "HTML解析库未安装，请安装beautifulsoup4",
                "file_type": "html",
                "content": None
            }
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'lxml')
            
            # 移除脚本和样式
            for script in soup(["script", "style"]):
                script.decompose()
            
            # 提取文本
            text = soup.get_text(separator='\n', strip=True)
            
            # 提取标题
            title = soup.title.string if soup.title else None
            
            # 提取链接
            links = []
            for a in soup.find_all('a', href=True):
                links.append({
                    "text": a.get_text().strip(),
                    "href": a['href']
                })
            
            # 清理文本
            text = self._clean_text(text)
            
            return {
                "success": True,
                "error": None,
                "file_type": "html",
                "content": text,
                "metadata": {
                    "title": title,
                    "links": links,
                    "char_count": len(text)
                }
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"HTML解析失败: {str(e)}",
                "file_type": "html",
                "content": None
            }
    
    def _clean_text(self, text: str) -> str:
        """清理文本内容"""
        if not text:
            return ""
        
        # 替换多个空格为单个空格
        text = re.sub(r' +', ' ', text)
        
        # 替换多个换行为两个换行
        text = re.sub(r'\n\n+', '\n\n', text)
        
        # 移除多余的空白字符
        text = re.sub(r'[ \t]+', ' ', text)
        
        # 去除首尾空白
        text = text.strip()
        
        # 限制文本长度
        if len(text) > self.max_text_length:
            text = text[:self.max_text_length] + "...[文本过长，已截断]"
        
        return text
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """提取文件元数据"""
        try:
            stat = os.stat(file_path)
            
            metadata = {
                "file_name": os.path.basename(file_path),
                "file_size": stat.st_size,
                "created_time": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                "modified_time": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                "file_type": self.get_file_type(file_path)
            }
            
            return metadata
            
        except Exception as e:
            logger.error(f"元数据提取失败: {e}")
            return {
                "file_name": os.path.basename(file_path),
                "file_size": 0,
                "created_time": None,
                "modified_time": None,
                "file_type": None,
                "error": str(e)
            }
    
    def validate_file(self, file_path: str) -> Dict[str, Any]:
        """验证文件是否可解析"""
        try:
            if not os.path.exists(file_path):
                return {
                    "valid": False,
                    "error": "文件不存在"
                }
            
            file_type = self.get_file_type(file_path)
            if file_type is None:
                return {
                    "valid": False,
                    "error": "不支持的文件类型"
                }
            
            file_size = os.path.getsize(file_path)
            if file_size > self.max_file_size:
                return {
                    "valid": False,
                    "error": f"文件大小超过限制（最大{self.max_file_size // 1024 // 1024}MB）"
                }
            
            return {
                "valid": True,
                "file_type": file_type,
                "file_size": file_size
            }
            
        except Exception as e:
            return {
                "valid": False,
                "error": f"文件验证失败: {str(e)}"
            }


# 全局解析器实例
file_parser = FileParser()

# 便捷函数
def parse_file(file_path: str, file_type: Optional[str] = None) -> Dict[str, Any]:
    """解析文件的便捷函数"""
    return file_parser.parse_file(file_path, file_type)

def get_file_type(file_path: str) -> Optional[str]:
    """获取文件类型的便捷函数"""
    return file_parser.get_file_type(file_path)

def validate_file(file_path: str) -> Dict[str, Any]:
    """验证文件的便捷函数"""
    return file_parser.validate_file(file_path)

def extract_metadata(file_path: str) -> Dict[str, Any]:
    """提取文件元数据的便捷函数"""
    return file_parser.extract_metadata(file_path)
